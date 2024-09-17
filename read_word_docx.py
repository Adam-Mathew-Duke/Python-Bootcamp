# Reading Microsoft Word Documents

# add on module is called python-docx
# you can just say import docx
# this is for Python version 3
import docx

doc = docx.Document('testdata.docx')

# get the paragraphs from the document
doc.paragraphs

# get the number of lines in the document
len(doc.paragraphs)

# print friendly version
# print the first line of text from our document
doc.paragraphs[0].text

# print all our Word Documetns lines to the terminal
# put the character count at the end of each line
for paragraph in doc.paragraphs:
    print(paragraph.text+' ('+str(len(paragraph.text))+' chars)')

# end of code
